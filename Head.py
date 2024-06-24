# -*- coding: utf-8 -*-
"""
Created on Mon Jul 17 11:11:07 2023

@author: Mohamed Mustafa
"""
import numpy as np
import time
import pandas as pd
import sys
#from fuzzywuzzy import fuzz
import docx
import sqlite3 as db


import matplotlib.pyplot as plt
import matplotlib.image as mpimg


time_limit = 300
start_time=time.time()
end_time=start_time+time_limit
print("Hai, My name is Ria, your diabetes assistant!")
print("I'm here to answer your doubts, Let's start!")
name = input("What's your name?")
print('Hi '+name+' How are you?')
s = input()

'''mobile_number = input("Enter your mobile number: ")
email = input("Enter your email id: ")

a = np.array([name])
b = np.array([mobile_number])
c = np.array([email])
frame={"Name":a,"Phone no":b,"Mail id":c}
df = pd.DataFrame(frame)
database = "DiabetesBot.sqlite"
conn = db.connect(database)
df.to_sql(name='Users', con=conn,if_exists='append',index=False)
conn.close()'''

def readtxt(filename):
    doc = docx.Document(filename)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    return '\n'.join(fullText)

def moreQueries():
    a = input('Do you have any queries? Y|N: ')
    if(a.lower()=='y'):
        main()
    elif(a.lower()=='n'):
        print('Thank you, Have a nice day!')
        sys.exit()
    else:
        print('Invalid option')
        sys.exit()
        
def is_valid_mobile_number(mobile_number):
    if len(mobile_number) == 10 and mobile_number.isdigit():
        return True
    else:
        return False

def is_valid_email(email):
    if '@' in email and '.' in email:
        return True
    else:
        return False


'''while not is_valid_mobile_number(mobile_number):
    print("Invalid mobile number. Please enter a 10-digit number.")
    mobile_number = input("\nEnter a mobile number: ")

#Validate the email
while not is_valid_email(email):
    print("Invalid email address. Please enter a valid email.")
    email = input("\nEnter your email id: ")'''
    
def dailyLife():
    print('Select any one of these:(press the given options as follows)\n1.Does Walking reduces my blood glucose level?\n2.Can I skip breakfast?\n3.What would be right time to have my dinner daily?\n4.Can I have any detox drink?')
    removeWords = ['what','waht','wat','is','a','how','the','?', 'of','are','tell','me','name','that','to','can','be','it','?','any','other','there','followed','may','for','required','does','do','did','should','shall','i','my','which','can','could','have','had','has','right','show','would','will','daily','everyday','patients','patient']
    a = input()
    Words = a.split()
    filteredWords1 = [word for word in Words if word.lower() not in removeWords]
  #  print(filteredWords)
    sentence1 = " ".join(filteredWords1)   
    z = sentence1.lower()
    z = z.replace('?',' ')
    print(z)
    return z

def dailyOptions(b):
    if b==str(1) or b in ['walking','walking reduce blood glucose level','walking reduces blood glucose level','walking reduce blood sugar level']:
        
        print("yes, absolutely!  a 30-minute brisk walk within 30 minutes after a meal can lower your blood sugar 50 times more than being sedentary")  
    elif b==str(2) or b in ['skip breakfast','skips breakfast']:
         print("It's commonly said that breakfast is the most important meal of the day, skipping breakfast occasionally can even raise your risk.\nEating Breakfast every day can lower your risk for Type 2 Diabetes")
    elif b==str(3) or b in ['time dinner','right time dinner']:
         print("Dinner should be about 3 hours prior to sleeping. If you sleep at around 11 at night, ideal dinner time for you is between 8-8.30. -if you sleep at 10 pm, then ideal dinner time for you is 7-7.30 pm; so you can accommodate a bed time snack.")
        
    elif b==str(4) or b in ['detox drink']:
        detoxOptions(b)
        y = input("\nPRESS 1 --> to go BACK (or) PRESS 2 --> MAIN menu:\n")
        while y==str(1):
            detoxOptions(b)
            y = input("\nPRESS 1 --> to go BACK (or) PRESS 2 --> MAIN menu:\n")
        else:
            main()
        

def detoxOptions(x):
    print("yes, you should definitely add to your diet.\nHere are 5 nutritious detox diabetic drinks to reduce high blood sugar levels\n1.Tulsi water\n2.Ginger water\n3.Methi water\n4.Cinnamon water\n5.Neem water")
    x = input("Are you aware of these detox drinks? if you want to know the recipes enter the above option or detox drink name:\n")
    if x==str(1) or x.lower() in ["tulsi water",'tulasi water', 'tulsiwater']:
        print("Tulsi, also known as basil is loaded with hypoglycaemic properties that help to maintain proper blood sugar levels in the body. You need to boil 6-8 tulsi in a glass of water and drink it hot or cold whenever it is possible for you throughout the day")
    elif x==str(2) or x.lower() in ['ginger water','ginger','gingerwater']:
        print('It is a no-brainer that ginger has many health benefits. It contains zinc that promotes the secretion of insulin. Moreover, it is also jam-packed with antioxidants and anti-inflammatory properties. Hence, boiling ginger roots in a glass of water and straining and drinking it can be beneficial for diabetics')
    elif x==str(3) or x.lower() in ['methi water','methi','methiwater','fenugreek water','fenugreek']:
       print('Methi (fenugreek) is helpful in dealing with insulin resistance. It is a good option for those with diabetes. Try to soak methi overnight in water, and drink it the next morning. Remember to boil and strain it before drinking the water')
    elif x==str(4) or x.lower() in ['cinnamon water','cinnamon','cinnamonwater']:
        print('Cinnamon helps the pancreas to release insulin that tends to promote glucose processi4ng in the body. So, soak 1 teaspoon of cinnamon powder in a glass of water overnight and drink it the next morning. You will surely be able to manage your blood sugar levels')
    elif x==str(5) or x.lower() in ['neem water','neem','neemwater']:
        print('It can do wonders to the health. Yes, you have heard it right! A majority of people avoid opting for neem. But, it can be a boon for those having diabetes. Neem leaves contain anti-inflammatory and antiviral properties that can help diabetics to maintain blood sugar levels in the recommended range. Boil 7-8 neem leaves in a glass of water and drink it. Though the taste will be bitter and pungent, it is good for your overall well-being.')    
        
def preventionOptions():
    a = input('press any given option to know more about this or press 0 -->to Main menu \n')
    if a==str(1) or a.lower() in ['losing weight','lose weight']:
        print('Losing weight reduces the risk of diabetes. People in one large study reduced their risk of developing diabetes by almost 60% after losing approximately 7% of their body weight with changes in exercise and diet.')
        print('Set a weight-loss goal based on your current body weight. Talk to your doctor about reasonable short-term goals and expectations, such as a losing 1 to 2 pounds a week.')
       
    elif a==str(2) or a.lower() in ['healthy diet','healthy eating plan','healthy food']:
        with open("Prevention-healthyDiet.txt", "r") as file:
             file_contents = file.read()
             print(file_contents)
    elif a==str(3) or a.lower() in ['exercise','regular exercise','get regular exercise']:
        print('There are many benefits to regular physical activity. Exercise can help you:\n*Lose weight\n*Lower your blood sugar\n*Boost your sensitivity to insulin — which helps keep your blood sugar within a normal range\nSet a weight-loss goal based on your current body weight. Talk to your doctor about reasonable short-term goals and expectations, such as a losing 1 to 2 pounds a week.\nGoals for most adults to promote weight loss and maintain a healthy weight include:\n   *Aerobic exercise. Aim for 30 minutes or more of moderate to vigorous aerobic exercise — such as brisk walking, swimming, biking or running — on most days for a total of at least 150 minutes a week.Resistance exercise.\n   *Resistance exercise — at least 2 to 3 times a week — increases your strength, balance and ability to maintain an active life. Resistance training includes weightlifting, yoga and calisthenics.\n   *Limited inactivity. Breaking up long bouts of inactivity, such as sitting at the computer, can help control blood sugar levels. Take a few minutes to stand, walk around or do some light activity every 30 minutes.')
   
    elif a == str(4) or a.lower()==['avoid alcohol','avoid smoking','don;t smoke']:
        print('Smoking can contribute to insulin resistance, which can lead to type 2 diabetes. If you already smoke, try to quit.')
    elif a == str(5) or a.lower() in ['care provider','health care provider']:
        print('Share your concerns about diabetes prevention with your doctor. He or she will appreciate your efforts to prevent diabetes and may offer additional suggestions based on your medical history or other factors.\n If you are at high risk, your provider may suggest that you take one of a few types of diabetes medicines')
        
def read_para(filename,index):
     doc = docx.Document(filename)
     para = doc.paragraphs[index]
     #print('len of para',len(doc.paragraphs))
     print(para.text)

def main():
    while time.time()<end_time:
    
        time.sleep(1)
        print('\nGood, Ask me any question or Enter the given option:\n')
        image_filenames = ["diabetes.webp","treatmentImage.png", "diet.jpeg","dailylife.webp","preventions.jpeg"]
        texts = ["1.Diabetes and diagnosis", "2.Treatment", "3.Diet","4.Daily-life","5.Prevention"]
        fig, axs = plt.subplots(5, 2, figsize=(4,2))  # 3 rows, 2 columns
        for i in range(len(image_filenames)):
            img = mpimg.imread(image_filenames[i])
            row = i
            axs[row, 1].imshow(img)
            axs[row, 1].axis('off')
            
            axs[row, 0].text(0.5, 0.5, texts[i], fontsize=12, ha='left', va='center')
            axs[row, 0].axis('off') 
        plt.tight_layout()
        plt.show()
        
        
      
       
        ques = input()
        ques = ques.replace('?', ' ')
        removeWords = ['what','waht','wat','is','a','how','the','?', 'of','are','tell','me','name','that','to','can','be','it','?','any','other','there','followed','may','for','required','does','do','did','should','shall','i','my','which','can','could','have','had','has','right','show','will','would','daily','everyday']
        words = ques.split()
        filteredWords = [word for word in words if word.lower() not in removeWords]
      #  print(filteredWords)
        sentence = " ".join(filteredWords)   
      #  print("sentence: ", sentence.lower())
       
       
        c = sentence.lower()
        
        
        #print(c)
        if (c in ['diabete','diabetes','blood sugar','blood glucose'] or c == str(1)):
             with open("diabetes.txt", "r") as file:
                 file_contents = file.read()
                 print(file_contents)
                 moreQueries()
             
        elif (c in ['types diabetes','kinds','diabetes types']):
            print('There are three main types of diabetes:\n1.Type 1 diabetes(Insulin-dependent diabetes mellitus)\n2.Type 2 diabetes(Non-Insulin dependent diabetes mellitus)\n3.Gestational diabetes (diabetes while pregnant)')
            a = input('Do you want to know more about these types? If so then Press 1\n')
            if a == str(1) or a.lower() in ['yes','y','more']:                      
                print("\nType1 diabetes: \n  In diabetes type 1, the pancreas does not make insulin, because the body's immune system attacks the islet cells in the pancreas that make insulin\n\nType2 diabetes: \n  In diabetes type 2, the pancreas makes less insulin than used to, and your body becomes resistant to insulin.\n\nGestational diabetes:\n  During pregnancy, some people may develop high blood sugar levels. This condition is known as gestational diabetes mellitus (GDM) or gestational diabetes. Gestational diabetes typically develops between the 24th and 28th weeks of pregnancy.")
              #  moreQueries()
            else:
               moreQueries()
               
        elif c in ['gestational diabetes','pregnancy diabetes','pregnancy sugar','gestational blood sugar','pregnancy blood sugar']:
            print("During pregnancy, some people may develop high blood sugar levels. This condition is known as gestational diabetes mellitus (GDM) or gestational diabetes. Gestational diabetes typically develops between the 24th and 28th weeks of pregnancy.")
        elif c in ['type 1 diabetes','type1 diabetes','diabetes type1','diabetes type 1']:
             print("In diabetes type 1, the pancreas does not make insulin, because the body's immune system attacks the islet cells in the pancreas that make insulin\n")
        elif c in ['type 2 diabetes','type2 diabetes','diabetes type2','diabetes type 2']:
            print("In diabetes type 2, the pancreas makes less insulin than used to, and your body becomes resistant to insulin.")
        elif c in ['diabetes and diagnosis']:
             with open("diabetes.txt", "r") as file:
                 file_contents = file.read()
                 print(file_contents)
                 moreQueries()
                              
            # print(df)
        elif c in ['diagnosis','diagnosed', 'tested','sugar test','test','diagnose']  :
             with open("DiabetesTest.txt", "r") as file:
                file_contents = file.read()
                print(file_contents) 
                moreQueries()
                
        elif c==str(2) or c in ['treatment','treatment diabetes','treatement diabetes','treated','treacted']:
            with open("Treatment.txt", "r") as file:
                file_contents = file.read()
                print(file_contents) 
                a = input()
                if a==str(1):
                    print("If you have type 1 diabetes, you'll need to take insulin shots (or wear an insulin pump) every day. Insulin is needed to manage your blood sugar levels and give your body energy. You can't take insulin as a pill. That's because the acid in your stomach would destroy it before it could get into your bloodstream.")
                    moreQueries()
                elif a==str(2):
                    print("It is often treated with diet, exercise, medication taken orally, or with direct infusions of insulin. Oral medications help the insulin the body produces to be more effective in processing glucose. Insulin injected through a syringe or delivered by an insulin pump mimics the way the pancreas would naturally produce and distribute insulin")
                    moreQueries()
                elif a==str(3):
                    with open("Treatment-gestational.txt", "r") as file:
                        file_contents = file.read()
                        print(file_contents)
                        moreQueries()
                elif a==str(4):
                    print("1.Eat healthy foods.\n2.Be more active.\n3.Lose excess weight.\n4.Take medications as needed.")
                    moreQueries()
                elif a==str(4) or a.lower() in['medicines','common medications','medications','common medicines']:
                    with open("Medications.txt", "r") as file:
                        file_contents = file.read()
                        print(file_contents) 
                
        elif  c in ['treatment type1','type1 treatment','type1 treated','treated type1','type1 diabetes treated','type1 diabetes treatment','treat type1 diabetes','treated type 1','type 1 diabetes treated','type 1 diabetes treatment','treat type1 diabetes']:
            with open("TreatmentType1.txt", "r") as file:
                file_contents = file.read()
                print(file_contents) 
                moreQueries()
        
        elif  c in ['treat type2 diabetes','treatment type2','type2 treatment','type2 treated','treated type2','type2 diabetes treated','type2 diabetes treatment','type2','type 2','type 2 diabetes treated','typeii diabetes treatment','type 2 diabetes treatment']:
            with open("TreatmentType2.txt", "r") as file:
                file_contents = file.read()
                print(file_contents) 
                moreQueries()
                
            
        elif c==str(3) or c in ['diet','diet diabetic','diet diabetes','diat diabetes','diat diabetic']:
           with open("Diet.txt", "r") as file:
               file_contents = file.read()
               print(file_contents) 
               moreQueries()
                
               
               
        elif c==str(4) or c in ['daily-life','dailylife','daily life']:
            #  with open("diet.docx", "r") as file:
                
                  b = dailyLife()
                  dailyOptions(b)
                      #moreQueries()
                  y = input("\nPRESS 1 --> to go BACK or PRESS 2 --> MAIN menu:\n")
                  while y==str(1):
                      b = dailyLife()
                      dailyOptions(b)
                          
                  else:
                      main()
                      
                
        elif c in ['hypoglycemia','hypo glycemia','hypo glycemia']:
             print("Hypoglycemia is a condition in which your blood sugar (glucose) level is lower than the standard range. Glucose is your body's main energy source.\nHypoglycemia is often related to diabetes treatment")
                          
        elif c in ['treatment hypoglycemia','hypoglycemia treatment']:
             print('Treatment involves quickly getting your blood sugar back to within the standard range either with a high-sugar food or drink or with medication. Long-term treatment requires identifying and treating the cause of hypoglycemia')
        elif c in ['symptoms hypoglycemia']:
            with open("Symptoms-hypogycemia.txt", "r") as file:
                 file_contents = file.read()
                 print(file_contents)
        elif c in ['causes hypoglycemia']:
             with open("causes-hypogycemia.txt", "r") as file:
                  file_contents = file.read()
                  print(file_contents)
        elif c in ['hyperglcemia']:
            print('Hyperglycemia happens when there’s too much sugar (glucose) in your blood. It’s also called high blood sugar or high blood glucose. This happens when your body has too little insulin (a hormone) or if your body can’t use insulin properly (insulin resistance).')
           
        elif c in ['symptoms diabetes']:
             with open("Symptoms-diabetes.txt", "r") as file:
                  file_contents = file.read()
                  print(file_contents)
                 
        elif c in['symptoms type1 diabetes','symptoms type 1 symptoms','type1 diabetes symptoms','type 1 diabetes symptoms']:
            print('People who have type 1 diabetes may also have nausea, vomiting, or stomach pains. Type 1 diabetes can be diagnosed at any age, and symptoms can develop in just a few weeks or months and can be severe')
        
        elif c in['symptoms type2 diabetes','symptoms type 2 symptoms','type2 diabetes symptoms','type 2 diabetes symptoms']:
            print('Type 2 diabetes symptoms often take several years to develop. Some people don’t notice any symptoms at all. Type 2 diabetes usually starts when you’re an adult, though more and more children and teens are developing it. Because symptoms are hard to spot, it’s important to know the risk factors for type 2 diabetes. Make sure to visit your doctor if you have any of them.')    
            with open("risk factors-type2.txt", "r") as file:
                 file_contents = file.read()
                 print(file_contents)
        elif c in ['gi','glycemic index']:
             with open("GI.txt", "r") as file:
                  file_contents = file.read()
                  print(file_contents)
        elif c in ['symptoms gestational diabetes','gestational diabetes symptoms','pregnancy sugar symptoms','pregnancy blood sugar symptoms']:
            print('Gestational diabetes (diabetes during pregnancy) usually doesn’t have any symptoms. If you’re pregnant, your doctor should test you for gestational diabetes between 24 and 28 weeks of pregnancy. If needed, you can make changes to protect your health and your baby’s health.')
        elif c in ['risk factors type 1 diabetes']:
           print('Known risk factors include:\nFamily history: Having a parent, brother, or sister with type 1 diabetes.\nAge: You can get type 1 diabetes at any age, but it usually develops in children, teens, or young adults')
        elif c in ['risk factors type 2 diabetes']:
            with open("risk factors-type2.txt", "r") as file:
                 file_contents = file.read()
                 print(file_contents) 
            
        elif c in ['risk factors prediabetes']:
            with open("risk factors-pre.txt", "r") as file:
                 file_contents = file.read()
                 print(file_contents) 
                
        elif c in ['risk factors gestational diabetes']:
            with open("risk factors-gestational.txt", "r") as file:
                 file_contents = file.read()
                 print(file_contents) 
        elif c==str(5) or c in['diabetes preventions','diabetes prevented','blood sugar preventions','blood sugar prevented']:
           with open("preventions.txt", "r") as file:
                file_contents = file.read()
                print(file_contents) 
                preventionOptions()
                
                moreQueries()
             
main()       
        
        
        
        
        
        
        